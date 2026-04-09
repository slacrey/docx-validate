import re
import zipfile
from collections.abc import Iterable
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph


BOOKMARK_NAMESPACE = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
PLACEHOLDER_PATTERNS = (
    re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}"),
    re.compile(r"\$\{([A-Za-z0-9_]+)\}"),
)


def extract_template_metadata(path: str | Path) -> dict[str, object]:
    document_path = Path(path)
    document = Document(document_path)

    return {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "section_count": len(document.sections),
        "styles": _extract_styles(document),
        "bookmarks": _extract_bookmarks(document_path),
        "placeholders": _extract_placeholders(document),
    }


def _extract_styles(document: DocxDocument) -> list[str]:
    styles = {paragraph.style.name for paragraph in _iter_paragraphs(document)}
    return sorted(styles)


def _extract_bookmarks(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("word/document.xml"))

    bookmarks = {
        node.attrib[f"{{{BOOKMARK_NAMESPACE['w']}}}name"]
        for node in root.findall(".//w:bookmarkStart", BOOKMARK_NAMESPACE)
        if node.attrib.get(f"{{{BOOKMARK_NAMESPACE['w']}}}name")
        and not node.attrib[f"{{{BOOKMARK_NAMESPACE['w']}}}name"].startswith("_")
    }
    return sorted(bookmarks)


def _extract_placeholders(document: DocxDocument) -> list[str]:
    placeholders: set[str] = set()
    for paragraph in _iter_paragraphs(document):
        for pattern in PLACEHOLDER_PATTERNS:
            placeholders.update(pattern.findall(paragraph.text))
    return sorted(placeholders)


def _iter_paragraphs(document: DocxDocument) -> Iterable[Paragraph]:
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
